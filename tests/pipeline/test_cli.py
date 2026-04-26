from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]


def run_pipeline(args: list[str], *, cwd: Path = ROOT) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, "-m", "tools.pipeline", *args],
        cwd=cwd,
        capture_output=True,
        text=True,
        check=False,
    )


def test_draft_cli_generates_manifest_validation_docx_for_three_fixtures(tmp_path: Path) -> None:
    requests = [
        {
            "documentType": "corporate",
            "targetLanguage": "ko",
            "jurisdiction": "korea",
            "title": "이사회 결의서",
            "description": "board-resolution",
            "sections": [
                {"heading": "결의 사항", "body": "본 이사회는 신주 발행 준비 절차를 진행하기로 의결한다."},
                {"heading": "시행일", "body": "본 결의는 2026년 4월 26일부터 시행한다."},
            ],
        },
        {
            "documentType": "general",
            "targetLanguage": "en",
            "jurisdiction": "international",
            "title": "Compliance Note",
            "description": "compliance-note",
            "sections": [
                {"heading": "Purpose", "body": "This note records supplied compliance steps."},
                {"heading": "Checklist", "body": "| Item | Status |\n|---|---|\n| Notice | Complete |"},
            ],
        },
        {
            "documentType": "advisory",
            "supportLevel": "conditional",
            "authorityPacketProvided": True,
            "targetLanguage": "en",
            "jurisdiction": "us",
            "title": "Authority-Supported Memo",
            "description": "authority-memo",
            "sections": [
                {"heading": "Issue", "body": "This memo summarizes the user-provided authority packet."},
                {"heading": "Conclusion", "body": "The user-supplied conclusion is preserved for counsel review."},
            ],
        },
    ]

    for index, payload in enumerate(requests, start=1):
        request = tmp_path / f"request-{index}.json"
        output_base = tmp_path / f"out-{index}"
        request.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

        completed = run_pipeline(["draft", "--request", str(request), "--output-base", str(output_base)])

        assert completed.returncode == 0, completed.stderr + completed.stdout
        summary = json.loads(completed.stdout)
        assert summary["status"] == "completed"
        assert Path(summary["artifactPaths"]["manifest"]).exists()
        assert Path(summary["artifactPaths"]["validation_report"]).exists()
        assert Path(summary["artifactPaths"]["rendered_output"]).exists()
        assert Path(summary["checkpointPath"]).exists()
        assert summary["completedStep"] == "D6"
        assert summary["nextResumeStep"] is None


def test_draft_cli_failure_writes_checkpoint_with_resume_step(tmp_path: Path) -> None:
    request = tmp_path / "bad-request.json"
    output_base = tmp_path / "bad-out"
    request.write_text(
        json.dumps(
            {
                "documentType": "general",
                "targetLanguage": "ko",
                "jurisdiction": "korea",
                "title": "번호 오류 문서",
                "content": "# 번호 오류 문서\n\n제1조(목적) 본 문서는 테스트이다.\n제3조(시행일) 번호가 건너뛴다.\n",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    completed = run_pipeline(["draft", "--request", str(request), "--output-base", str(output_base)])

    assert completed.returncode == 1
    summary = json.loads(completed.stdout)
    checkpoint = json.loads(Path(summary["checkpointPath"]).read_text(encoding="utf-8"))
    assert summary["status"] == "failed"
    assert summary["failedStep"] == "D4"
    assert summary["nextResumeStep"] == "D3"
    assert checkpoint["status"] == "failed"
    assert checkpoint["nextResumeStep"] == "D3"
    assert "rendered_output" not in summary["artifactPaths"]


def test_revise_cli_generates_docx_and_level_b_artifacts_from_docx_input(tmp_path: Path) -> None:
    original = tmp_path / "original.docx"
    doc = Document()
    doc.add_heading("Original Opinion", level=1)
    doc.add_paragraph("This paragraph records the original supplied text.")
    doc.save(original)
    instructions = tmp_path / "instructions.md"
    instructions.write_text("Add a short note that the issue was revised per user instruction.", encoding="utf-8")
    output_base = tmp_path / "revise-out"

    completed = run_pipeline(
        [
            "revise",
            "--input",
            str(original),
            "--instructions",
            str(instructions),
            "--output-base",
            str(output_base),
        ]
    )

    assert completed.returncode == 0, completed.stderr + completed.stdout
    summary = json.loads(completed.stdout)
    assert summary["status"] == "completed"
    assert summary["completedStep"] == "R7"
    assert Path(summary["artifactPaths"]["rendered_output"]).exists()
    assert Path(summary["artifactPaths"]["redlineOutput"]).exists()
    assert Path(summary["artifactPaths"]["changeMap"]).exists()
    change_map = json.loads(Path(summary["artifactPaths"]["changeMap"]).read_text(encoding="utf-8"))
    assert change_map["trackingLevel"] == "level-b"
    assert change_map["nativeTrackedChanges"] is False

