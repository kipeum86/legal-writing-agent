from __future__ import annotations

import hashlib
import json
import subprocess
import sys
from pathlib import Path

from docx import Document

from tools.artifacts import schemas
from tools.parsing.docx_parser import parse_docx


ROOT = Path(__file__).resolve().parents[2]


def make_fixture(path: Path) -> None:
    doc = Document()
    doc.add_heading("법률검토의견", level=1)
    doc.add_paragraph("본 문서는 구조 추출 테스트를 위한 합성 문서입니다.")
    doc.add_heading("1. 검토의 배경", level=2)
    doc.add_paragraph("제공된 사실관계를 기준으로 검토 범위를 정리합니다.")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "자료"
    table.cell(1, 1).text = "합성 fixture"

    doc.add_heading("2. 검토의견", level=2)
    doc.add_paragraph("제1조(목적) 이 조항은 번호 후보로 인식되어야 합니다.")
    doc.add_paragraph("첫 번째 목록 항목", style="List Bullet")
    doc.save(path)


def test_parse_docx_extracts_blocks_headings_tables_and_hashes(tmp_path: Path) -> None:
    docx_path = tmp_path / "fixture.docx"
    make_fixture(docx_path)

    result = parse_docx(docx_path, document_id="doc-test")
    profile = result.profile

    assert profile["documentId"] == "doc-test"
    assert profile["sourceSha256"] == hashlib.sha256(docx_path.read_bytes()).hexdigest()
    assert profile["extractedTextSha256"]
    assert profile["blockCount"] >= 7

    headings = profile["headingTree"]
    assert [(node["title"], node["level"], node["parentId"]) for node in headings] == [
        ("법률검토의견", 1, None),
        ("1. 검토의 배경", 2, "s1"),
        ("2. 검토의견", 2, "s1"),
    ]

    tables = [block for block in profile["blocks"] if block["type"] == "table"]
    assert len(tables) == 1
    assert tables[0]["rows"] == [["항목", "내용"], ["자료", "합성 fixture"]]
    assert "합성 fixture" in tables[0]["text"]

    numbering = [
        block["numberingCandidate"]
        for block in profile["blocks"]
        if block.get("numberingCandidate")
    ]
    assert any(item["pattern"] == "korean_article" for item in numbering)
    assert any(item["isListStyle"] for item in numbering)


def test_parse_docx_builds_outline_and_clause_map_seed(tmp_path: Path) -> None:
    docx_path = tmp_path / "fixture.docx"
    make_fixture(docx_path)

    result = parse_docx(docx_path, document_id="doc-seed")
    outline = schemas.validate_artifact(result.outline, expected_type="outline")
    clause_map = schemas.validate_artifact(result.clause_map, expected_type="clause_map")

    assert outline["documentId"] == "doc-seed"
    assert clause_map["documentId"] == "doc-seed"
    assert outline["title"] == "법률검토의견"
    assert [section["sectionId"] for section in outline["sections"]] == ["s1", "s2", "s3"]
    assert [clause["sectionId"] for clause in clause_map["clauses"]] == ["s1", "s2", "s3"]
    assert all(clause["sourceTextHash"] for clause in clause_map["clauses"])
    assert clause_map["sourceProvenance"][0]["sha256"] == hashlib.sha256(docx_path.read_bytes()).hexdigest()


def test_docx_parser_cli_emits_combined_payload(tmp_path: Path) -> None:
    docx_path = tmp_path / "fixture.docx"
    out = tmp_path / "parsed.json"
    make_fixture(docx_path)

    completed = subprocess.run(
        [
            sys.executable,
            "-m",
            "tools.parsing.docx_parser",
            str(docx_path),
            "--document-id",
            "doc-cli",
            "--out",
            str(out),
        ],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )

    assert completed.returncode == 0
    stdout_payload = json.loads(completed.stdout)
    file_payload = json.loads(out.read_text(encoding="utf-8"))
    assert stdout_payload == file_payload
    assert stdout_payload["profile"]["documentId"] == "doc-cli"
    assert stdout_payload["outline"]["artifactType"] == "outline"
    assert stdout_payload["clauseMap"]["artifactType"] == "clause_map"
