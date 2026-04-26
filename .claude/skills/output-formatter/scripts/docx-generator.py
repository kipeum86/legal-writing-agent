#!/usr/bin/env python3
"""
docx-generator.py
Generates .docx files from markdown content with legal document formatting.

Usage:
    python3 docx-generator.py <input.md> <output.docx> [--lang ko|en]
                              [--jurisdiction korea|us|uk|intl]
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Mm, Pt
except ImportError as exc:
    print(f"Error: python-docx is required: {exc}", file=sys.stderr)
    sys.exit(1)


PAGE_CONFIGS = {
    "ko-korea": {
        "page_width": Mm(210),
        "page_height": Mm(297),
        "margins": {"top": Mm(20), "bottom": Mm(15), "left": Mm(20), "right": Mm(20)},
        "body_font_ascii": "바탕체",
        "body_font_east_asia": "바탕체",
        "heading_font_ascii": "맑은 고딕",
        "heading_font_east_asia": "맑은 고딕",
        "body_size": 12,
        "line_spacing": 1.6,
    },
    "ko-korea-opinion": {
        "page_width": Mm(210),
        "page_height": Mm(297),
        "margins": {"top": Mm(25.4), "bottom": Mm(25.4), "left": Mm(25.4), "right": Mm(25.4)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "맑은 고딕",
        "heading_font_ascii": "Times New Roman",
        "heading_font_east_asia": "맑은 고딕",
        "body_size": 11,
        "line_spacing": 1.15,
    },
    "en-us": {
        "page_width": Inches(8.5),
        "page_height": Inches(11),
        "margins": {"top": Inches(1), "bottom": Inches(1), "left": Inches(1), "right": Inches(1)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "Times New Roman",
        "heading_font_ascii": "Times New Roman",
        "heading_font_east_asia": "Times New Roman",
        "body_size": 12,
        "line_spacing": 1.5,
    },
    "en-uk": {
        "page_width": Mm(210),
        "page_height": Mm(297),
        "margins": {"top": Mm(25.4), "bottom": Mm(25.4), "left": Mm(25.4), "right": Mm(25.4)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "Times New Roman",
        "heading_font_ascii": "Arial",
        "heading_font_east_asia": "Arial",
        "body_size": 12,
        "line_spacing": 1.5,
    },
    "en-intl": {
        "page_width": Mm(210),
        "page_height": Mm(297),
        "margins": {"top": Mm(25), "bottom": Mm(25), "left": Mm(25), "right": Mm(25)},
        "body_font_ascii": "Times New Roman",
        "body_font_east_asia": "Times New Roman",
        "heading_font_ascii": "Arial",
        "heading_font_east_asia": "Arial",
        "body_size": 12,
        "line_spacing": 1.5,
    },
}


def get_config_key(lang: str, jurisdiction: str, document_type: str | None = None) -> str:
    if lang == "ko":
        opinion_types = {
            "advisory",
            "legal-opinion",
            "legal-review-opinion",
            "client-memorandum",
        }
        if document_type in opinion_types:
            return "ko-korea-opinion"
        return "ko-korea"
    if jurisdiction == "us":
        return "en-us"
    if jurisdiction == "uk":
        return "en-uk"
    return "en-intl"


def apply_run_font(
    run,
    ascii_font: str,
    east_asia_font: str,
    size_pt: int,
    *,
    bold: bool = False,
    italic: bool = False,
) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
        run._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)


def apply_style_font(style, ascii_font: str, east_asia_font: str, size_pt: int) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    if style._element.rPr is not None and style._element.rPr.rFonts is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
        style._element.rPr.rFonts.set(qn("w:ascii"), ascii_font)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), ascii_font)


@dataclass
class MarkdownBlock:
    type: str
    text: str = ""
    level: int | None = None
    rows: list[list[str]] = field(default_factory=list)
    items: list[str] = field(default_factory=list)
    info: str = ""
    label: str = ""


def configure_document(document: Document, config: dict, *, classification: str | None = None) -> None:
    section = document.sections[0]
    section.page_width = config["page_width"]
    section.page_height = config["page_height"]
    section.top_margin = config["margins"]["top"]
    section.bottom_margin = config["margins"]["bottom"]
    section.left_margin = config["margins"]["left"]
    section.right_margin = config["margins"]["right"]

    apply_style_font(
        document.styles["Normal"],
        config["body_font_ascii"],
        config["body_font_east_asia"],
        config["body_size"],
    )
    apply_style_font(
        document.styles["Heading 1"],
        config["heading_font_ascii"],
        config["heading_font_east_asia"],
        16,
    )
    apply_style_font(
        document.styles["Heading 2"],
        config["heading_font_ascii"],
        config["heading_font_east_asia"],
        14,
    )
    apply_style_font(
        document.styles["Heading 3"],
        config["heading_font_ascii"],
        config["heading_font_east_asia"],
        12,
    )
    configure_custom_styles(document, config)
    configure_header_footer(section, config, classification=classification)


def configure_custom_styles(document: Document, config: dict) -> None:
    quote_style = ensure_paragraph_style(document, "Legal Block Quote", base_style="Normal")
    quote_style.paragraph_format.left_indent = Inches(0.35)
    quote_style.paragraph_format.right_indent = Inches(0.15)
    quote_style.paragraph_format.space_before = Pt(6)
    quote_style.paragraph_format.space_after = Pt(6)
    quote_style.paragraph_format.line_spacing = 1.15
    apply_style_font(
        quote_style,
        config["body_font_ascii"],
        config["body_font_east_asia"],
        max(config["body_size"] - 1, 9),
    )

    statute_style = ensure_paragraph_style(document, "Legal Statute Block", base_style="Normal")
    statute_style.paragraph_format.left_indent = Inches(0.25)
    statute_style.paragraph_format.right_indent = Inches(0.15)
    statute_style.paragraph_format.space_before = Pt(6)
    statute_style.paragraph_format.space_after = Pt(6)
    statute_style.paragraph_format.line_spacing = 1.1
    apply_style_font(
        statute_style,
        config["body_font_ascii"],
        config["body_font_east_asia"],
        max(config["body_size"] - 1, 9),
    )

    footnote_style = ensure_paragraph_style(document, "Legal Footnote Candidate", base_style="Normal")
    footnote_style.paragraph_format.left_indent = Inches(0.25)
    footnote_style.paragraph_format.first_line_indent = Inches(-0.25)
    footnote_style.paragraph_format.space_after = Pt(3)
    footnote_style.paragraph_format.line_spacing = 1.0
    apply_style_font(
        footnote_style,
        config["body_font_ascii"],
        config["body_font_east_asia"],
        max(config["body_size"] - 2, 8),
    )


def ensure_paragraph_style(document: Document, name: str, *, base_style: str):
    try:
        return document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles[base_style]
        return style


def configure_header_footer(section, config: dict, *, classification: str | None = None) -> None:
    if classification:
        paragraph = section.header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(classification)
        apply_run_font(
            run,
            config["body_font_ascii"],
            config["body_font_east_asia"],
            max(config["body_size"] - 2, 8),
        )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    apply_run_font(
        run,
        config["body_font_ascii"],
        config["body_font_east_asia"],
        max(config["body_size"] - 2, 8),
    )
    add_field(footer.add_run(), "PAGE")


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def add_paragraph(document: Document, text: str, config: dict, *, heading_level: int | None = None) -> None:
    if heading_level == 1:
        paragraph = document.add_paragraph(style="Heading 1")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif heading_level == 2:
        paragraph = document.add_paragraph(style="Heading 2")
    elif heading_level == 3:
        paragraph = document.add_paragraph(style="Heading 3")
    else:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    size = config["body_size"]
    ascii_font = config["body_font_ascii"]
    east_asia_font = config["body_font_east_asia"]
    if heading_level == 1:
        size = 16
        ascii_font = config["heading_font_ascii"]
        east_asia_font = config["heading_font_east_asia"]
        add_inline_runs(paragraph, text, ascii_font, east_asia_font, size, bold_default=True)
    elif heading_level == 2:
        size = 14
        ascii_font = config["heading_font_ascii"]
        east_asia_font = config["heading_font_east_asia"]
        add_inline_runs(paragraph, text, ascii_font, east_asia_font, size, bold_default=True)
    elif heading_level == 3:
        size = 12
        ascii_font = config["heading_font_ascii"]
        east_asia_font = config["heading_font_east_asia"]
        add_inline_runs(paragraph, text, ascii_font, east_asia_font, size, bold_default=True)
    else:
        add_inline_runs(paragraph, text, ascii_font, east_asia_font, size)

    fmt = paragraph.paragraph_format
    fmt.line_spacing = config["line_spacing"]
    if heading_level == 1:
        fmt.space_after = Pt(10)
    elif heading_level == 2:
        fmt.space_before = Pt(10)
        fmt.space_after = Pt(6)
    elif heading_level == 3:
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(4)
    else:
        fmt.space_after = Pt(6)


INLINE_TOKEN_RE = re.compile(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_)")


def add_inline_runs(
    paragraph,
    text: str,
    ascii_font: str,
    east_asia_font: str,
    size_pt: int,
    *,
    bold_default: bool = False,
    italic_default: bool = False,
) -> None:
    cursor = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            apply_run_font(
                run,
                ascii_font,
                east_asia_font,
                size_pt,
                bold=bold_default,
                italic=italic_default,
            )

        token = match.group(0)
        bold = bold_default
        italic = italic_default
        if token.startswith("***") and token.endswith("***"):
            value = token[3:-3]
            bold = True
            italic = True
        elif token.startswith("**") and token.endswith("**"):
            value = token[2:-2]
            bold = True
        elif token.startswith("__") and token.endswith("__"):
            value = token[2:-2]
            bold = True
        else:
            value = token[1:-1]
            italic = True

        run = paragraph.add_run(value)
        apply_run_font(run, ascii_font, east_asia_font, size_pt, bold=bold, italic=italic)
        cursor = match.end()

    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        apply_run_font(
            run,
            ascii_font,
            east_asia_font,
            size_pt,
            bold=bold_default,
            italic=italic_default,
        )


def parse_markdown(content: str) -> list[MarkdownBlock]:
    lines = content.splitlines()
    blocks: list[MarkdownBlock] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            blocks.append(MarkdownBlock("paragraph", text=" ".join(line.strip() for line in paragraph_lines)))
            paragraph_lines.clear()

    index = 0
    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        fence_match = re.match(r"^```(\w+)?\s*$", stripped)
        if fence_match:
            flush_paragraph()
            info = (fence_match.group(1) or "").lower()
            body: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                body.append(lines[index].rstrip())
                index += 1
            if index < len(lines):
                index += 1
            block_type = "statute_block" if info in {"law", "statute", "quote", "citation"} else "code_block"
            blocks.append(MarkdownBlock(block_type, text="\n".join(body), info=info))
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    "heading",
                    level=min(len(heading_match.group(1)), 3),
                    text=heading_match.group(2).strip(),
                )
            )
            index += 1
            continue

        footnote_match = re.match(r"^\[\^([^\]]+)\]:\s*(.+)$", line)
        if footnote_match:
            flush_paragraph()
            blocks.append(
                MarkdownBlock(
                    "footnote",
                    label=footnote_match.group(1).strip(),
                    text=footnote_match.group(2).strip(),
                )
            )
            index += 1
            continue

        if is_table_start(lines, index):
            flush_paragraph()
            rows: list[list[str]] = [split_table_row(lines[index])]
            index += 2
            while index < len(lines) and is_table_row(lines[index]):
                rows.append(split_table_row(lines[index]))
                index += 1
            blocks.append(MarkdownBlock("table", rows=rows))
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[index]).rstrip())
                index += 1
            blocks.append(MarkdownBlock("blockquote", text="\n".join(quote_lines)))
            continue

        ordered_match = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        bullet_match = re.match(r"^\s*[-*+]\s+(.+)$", line)
        if ordered_match or bullet_match:
            flush_paragraph()
            block_type = "ordered_list" if ordered_match else "bullet_list"
            items: list[str] = []
            pattern = r"^\s*\d+[.)]\s+(.+)$" if ordered_match else r"^\s*[-*+]\s+(.+)$"
            while index < len(lines):
                item_match = re.match(pattern, lines[index])
                if not item_match:
                    break
                items.append(item_match.group(1).strip())
                index += 1
            blocks.append(MarkdownBlock(block_type, items=items))
            continue

        paragraph_lines.append(line)
        index += 1

    flush_paragraph()
    return blocks


def is_table_start(lines: list[str], index: int) -> bool:
    return (
        index + 1 < len(lines)
        and is_table_row(lines[index])
        and is_table_separator(lines[index + 1])
    )


def is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_table_separator(line: str) -> bool:
    if not is_table_row(line):
        return False
    cells = split_table_row(line)
    return bool(cells) and all(re.match(r"^:?-{3,}:?$", cell.strip()) for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def render_block(document: Document, block: MarkdownBlock, config: dict) -> None:
    if block.type == "heading":
        add_paragraph(document, block.text, config, heading_level=block.level or 3)
    elif block.type == "paragraph":
        add_paragraph(document, block.text, config)
    elif block.type == "blockquote":
        add_quote_block(document, block.text, config, style="Legal Block Quote")
    elif block.type in {"statute_block", "code_block"}:
        add_quote_block(document, block.text, config, style="Legal Statute Block")
    elif block.type == "ordered_list":
        add_list(document, block.items, config, ordered=True)
    elif block.type == "bullet_list":
        add_list(document, block.items, config, ordered=False)
    elif block.type == "table":
        add_table(document, block.rows, config)
    elif block.type == "footnote":
        add_footnote_candidate(document, block.label, block.text, config)


def add_quote_block(document: Document, text: str, config: dict, *, style: str) -> None:
    for line in text.splitlines() or [""]:
        paragraph = document.add_paragraph(style=style)
        add_inline_runs(
            paragraph,
            line,
            config["body_font_ascii"],
            config["body_font_east_asia"],
            max(config["body_size"] - 1, 9),
        )


def add_list(document: Document, items: list[str], config: dict, *, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        paragraph = document.add_paragraph(style=style)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = config["line_spacing"]
        add_inline_runs(
            paragraph,
            item,
            config["body_font_ascii"],
            config["body_font_east_asia"],
            config["body_size"],
        )


def add_table(document: Document, rows: list[list[str]], config: dict) -> None:
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            text = row[column_index] if column_index < len(row) else ""
            add_inline_runs(
                paragraph,
                text,
                config["body_font_ascii"],
                config["body_font_east_asia"],
                max(config["body_size"] - 1, 9),
                bold_default=row_index == 0,
            )
            paragraph.paragraph_format.space_after = Pt(0)

    document.add_paragraph("")


def add_footnote_candidate(document: Document, label: str, text: str, config: dict) -> None:
    paragraph = document.add_paragraph(style="Legal Footnote Candidate")
    label_run = paragraph.add_run(f"{label}. ")
    apply_run_font(
        label_run,
        config["body_font_ascii"],
        config["body_font_east_asia"],
        max(config["body_size"] - 2, 8),
        bold=True,
    )
    add_inline_runs(
        paragraph,
        text,
        config["body_font_ascii"],
        config["body_font_east_asia"],
        max(config["body_size"] - 2, 8),
    )


def generate_docx(
    input_path: str,
    output_path: str,
    lang: str,
    jurisdiction: str,
    document_type: str | None = None,
    classification: str | None = None,
) -> None:
    input_file = Path(input_path)
    if not input_file.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    output_file = Path(output_path)
    output_file.parent.mkdir(parents=True, exist_ok=True)

    config = PAGE_CONFIGS[get_config_key(lang, jurisdiction, document_type)]
    content = input_file.read_text(encoding="utf-8")

    document = Document()
    configure_document(document, config, classification=classification)

    for block in parse_markdown(content):
        render_block(document, block, config)

    document.save(str(output_file))
    print(f"Document generated: {output_file}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a DOCX file from markdown using legal-document defaults."
    )
    parser.add_argument("input", help="Path to the input markdown file")
    parser.add_argument("output", help="Path to the output DOCX file")
    parser.add_argument("--lang", choices=["ko", "en"], default="ko")
    parser.add_argument(
        "--jurisdiction",
        choices=["korea", "us", "uk", "intl"],
        default="korea",
    )
    parser.add_argument(
        "--document-type",
        default=None,
        help="Optional document type hint (e.g. advisory, litigation, corporate)",
    )
    parser.add_argument(
        "--classification",
        default=None,
        help="Optional classification marker to render in the DOCX header",
    )
    args = parser.parse_args()

    try:
        generate_docx(
            args.input,
            args.output,
            args.lang,
            args.jurisdiction,
            args.document_type,
            args.classification,
        )
    except FileNotFoundError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        sys.exit(1)
    except Exception as exc:
        print(f"Error generating document: {exc}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
